from pathlib import Path

from docx import Document
from pypdf import PdfReader


def read_pdf(file_path: Path) -> str:
    """Extract text from a PDF resume."""

    reader = PdfReader(file_path)

    text = ""

    for page in reader.pages:
        page_text = page.extract_text()

        if page_text:
            text += page_text + "\n"

    return text


def read_docx(file_path: Path) -> str:
    """Extract text from a DOCX resume."""

    doc = Document(file_path)

    text = ""

    # Read normal paragraphs
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    # Read tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"

    return text


def read_resume(file_path: Path) -> str:
    """Read a PDF or DOCX resume and return its text."""

    extension = file_path.suffix.lower()

    if extension == ".pdf":
        return read_pdf(file_path)

    if extension == ".docx":
        return read_docx(file_path)

    raise ValueError(
        "Unsupported file format. Please provide a PDF or DOCX file."
    )