from email.mime import message
import json
import os
from pathlib import Path
from pydoc import text
from pydantic import BaseModel
import time
from turtle import st
from dotenv import load_dotenv
from groq import Groq
load_dotenv()
my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("API NOT FOUND")
Client = Groq(api_key=my_api_key)
model = "openai/gpt-oss-120b"
job_description = """
Job description
Job Title: AI Developer

Location: Indore (On-site)
Experience: 6 Months – 2 Years
Employment Type: Full-time

About the Role

We are looking for a motivated AI Developer to design, develop, and deploy AI-powered applications. The ideal candidate should have experience in Python, Machine Learning, Generative AI, and API integration, with a passion for building intelligent solutions using the latest AI technologies.

Key Responsibilities
• Develop and deploy AI/ML models for real-world applications.
• Build AI-powered web applications and automation solutions.
• Integrate Large Language Models (LLMs) such as OpenAI, Gemini, Claude, or Llama into applications.
• Develop AI chatbots, virtual assistants, and recommendation systems.
• Work with NLP, Computer Vision, or Predictive Analytics projects.
• Design and integrate REST APIs and AI services.
• Perform data preprocessing, feature engineering, and model evaluation.
• Optimize AI models for performance, scalability, and accuracy.
• Collaborate with developers, designers, and project managers to deliver AI solutions.
• Stay updated with the latest AI tools, frameworks, and industry trends.

Required Skills
• Strong proficiency in Python.
• Good understanding of Machine Learning and Deep Learning concepts.
• Experience with TensorFlow, PyTorch, or Scikit-learn.
• Knowledge of Generative AI, LLMs, Prompt Engineering, and RAG.
• Experience with AI APIs (OpenAI, Gemini, Claude, etc.).
• Familiarity with LangChain, LlamaIndex, or similar AI frameworks.
• Knowledge of SQL/NoSQL databases.
• Experience with Git and version control.
• Understanding of REST APIs and cloud deployment is a plus.

Preferred Qualifications
• Bachelor's degree in Computer Science, Information Technology, AI, Data Science, or a related field.
• 6 months to 2 years of relevant experience in AI/ML development.
• Experience working on AI-based projects or internships.
• Strong problem-solving and analytical skills

Pay: ₹10,000.00 - ₹30,000.00 per month

Work Location: In person"""

class JobD(BaseModel):
    role : str
    required_skills : list[str]
    preffered_skills : list[str]
    minimum_experience : float
    educational_requirement : list[str]
    responsibilities : list[str]
jobd_schema = JobD.model_json_schema()

system_prompt = f"""
You are a HR assistance.

You job is to analyze job description and the structured information from them.

Return ONLY valid JSON object that matches the following schema:
{jobd_schema}
IMPORTANT:
DO NOT return schema itself.
DO NOT fields like "properties", "title", or "type".
Fill the schema with the information extracted from the job description.
Do Not invent information that is not present in the job description.
"""
user_prompt = f""" Analyze the following job description and extract the required information:
{job_description} """
message_system = {"role": "system", "content": system_prompt}
message_user = {"role": "user", "content": user_prompt}
respose_format = {
    "type" : "json_object"
}
messages = [message_system, message_user]
response = Client.chat.completions.create(model=model, messages=messages, response_format=respose_format)

answer = response.choices[0].message.content
raw_json = answer


import json
job_data = json.loads(raw_json)
job = JobD(**job_data)

print(job.minimum_experience)
print(job.educational_requirement)

#real parse
class MatchResult(BaseModel):
    score : float 
    details : dict

class Experience(BaseModel):
    company : str | None=None
    role : str | None=None
    duration : str | None=None
    description : str | None=None
    skill_used : list[str]

class Resume(BaseModel):
    name : str | None=None
    email : str | None=None
    phone : str | None=None
    linkedin : str | None=None
    github : str | None=None
    total_experience : float | None=None
    education : list[str]
    skills : list[str]
    experiences : list[Experience]
    projects : list[str] 
    certifications : list[str] 

resume_schema = Resume.model_json_schema()

def final_score(job,resume):
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
    You are an HR recruiter.

    Compare the candidate's resume with the job description.

    JOB DESCRIPTION:
    {job.model_dump_json(indent=2)}

    CANDIDATE RESUME:
    {resume.model_dump_json(indent=2)}
    Return JSON matching this schema:

    {match_schema}

    Give me:

    1. Candidate name
    2. Matching skills
    3. Missing important skills
    4. Whether experience requirement is met
    5. Overall match percentage from 0 to 100
    6. A short final verdict

    Keep the response concise and easy to read.
    """
    message={
        "role": "user",
        "content" : prompt
    }
    messages=[message]
    response_format={
        "type": "json_object"
    }
    response = Client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    data = json.loads(response.choices[0].message.content)
    return MatchResult(**data)
def parse_resume(resume_text):
    system_prompt = f"""
    You are an expert resume parser.

    Extract information from the resume based on its meaning,
    not only based on exact section headings.

    Different resumes may use different headings.

    For example:
    - Experience
    - Professional Experience
    - Work History
    - Employment
    - Internships

    These may all contain relevant experience.

    Skills may also appear in the skills section, work experience,
    internships or projects.

    Return ONLY valid JSON matching this schema:

    {resume_schema}

    Important rules:

    1. Do not invent information.
    2. If a value is not available, return null.
    3. If a list has no information, return an empty list.
    4. Include internships inside experiences.
    5. Extract skills mentioned across the entire resume.
    """
    user_prompt = f"""
    Parse the following resume:

    {resume_text}
    """
    message_system={
        "role" : "system",
        "content" : system_prompt
    }
    message_user={
        "role" : "user",
        "content" : user_prompt
    }
    messages=[message_system, message_user]
    response_format={
        "type": "json_object"
    }
    response=Client.chat.completions.create(model=model, messages=messages, response_format=response_format)
    raw_output = response.choices[0].message.content
    data = json.loads(raw_output)
    resume = Resume(**data)
    return resume       



from pypdf import PdfReader
from docx import Document

def read_pdf(file_path):
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def read_docx(file_path):
    doc = Document(file_path)
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():  # Check if the cell is not empty
                    text += cell.text + "\n"    
    return text
def read_resume(file_path):
    if file_path.suffix.lower() == '.pdf':
        return read_pdf(file_path)
    elif file_path.suffix.lower() == '.docx':
        return read_docx(file_path)
    else:
        raise ValueError("Unsupported file format. Please provide a PDF or DOCX file.")

resume_folder = Path("Resumes")
all_results=[]
for file_path in resume_folder.iterdir():
    if not file_path.is_file():
        continue
    if file_path.suffix.lower() not in ['.pdf', '.docx']:
        continue
    resume_text = read_resume(file_path)
    parsed_resume = parse_resume(resume_text)
    time.sleep(5)  # Add a delay of 5 seconds between requests
    result = final_score(job, parsed_resume)
    time.sleep(5)  # Add a delay of 5 seconds between requests
    print("Score:", result.score)
    all_results.append({
        "file_name": file_path.name,
        "score": result.score,
        "details": result.details
    })
all_results.sort(
    key=lambda x: x['score'],
    reverse=True
    )
top_2 = all_results[:2]
worse_2 = all_results[-2:]

print(" TOP 2 CANDIDATES")
for candidate in top_2:
    print(f"File: {candidate['file_name']}, Score: {candidate['score']}")
    print("Details:", candidate['details'])

print(" WORST 2 CANDIDATES")
for candidate in worse_2:
    print(f"File: {candidate['file_name']}, Score: {candidate['score']}")
    print("Details:", candidate['details'])
